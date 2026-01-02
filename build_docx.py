#!/usr/bin/env python3
"""
ATS-Friendly Resume Builder
Generates role-specific resumes in DOCX format from YAML data
"""

import yaml
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(doc):
    """
    Add a horizontal line (section divider) to the document
    """
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    
    # Add bottom border to create a horizontal line
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Line thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')  # Black color
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # ZERO spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 0.01  # Almost no line height
    
    return para


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph with underline formatting
    """
    # Get the part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Georgia')
    rFonts.set(qn('w:hAnsi'), 'Georgia')
    rPr.append(rFonts)
    
    # Set font size (7.5pt = 15 half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '15')
    rPr.append(sz)
    
    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Set color to blue (optional, for visual hyperlink)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink


def filter_by_tags(items, role):
    """
    Filter items based on role tags
    """
    if role == "all":
        return items
    
    filtered = []
    for item in items:
        tags = item.get("tags", [])
        if "all" in tags or role in tags:
            filtered.append(item)
    return filtered


def apply_bold_to_text(paragraph, text):
    """
    Apply bold formatting to text within ** markers
    Example: "Increased by **27%**" -> "Increased by 27%" with 27% bolded
    """
    import re
    
    # Split text by bold markers
    parts = re.split(r'\*\*(.*?)\*\*', text)
    
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Georgia'
        run.font.size = Pt(7.5)
        
        # Odd indices are the bolded parts
        if i % 2 == 1:
            run.bold = True


def build_resume(role):
    """
    Main function to build resume DOCX for a specific role
    """
    
    # Load YAML data
    with open("data/resume.yaml", "r") as f:
        data = yaml.safe_load(f)
    
    # Create document
    doc = Document()
    
    # ===========================
    # DOCUMENT MARGINS (0.4 inch)
    # ===========================
    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    
    # ===========================
    # HEADER SECTION
    # ===========================
    personal = data["personal"]
    
    # Name and Title (Bold, 12pt, Georgia, CENTER ALIGNED)
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = header.add_run(f"{personal['name']} | {personal['title']}")
    name_run.bold = True
    name_run.font.name = 'Georgia'
    name_run.font.size = Pt(12)
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.space_before = Pt(0)
    
    # Contact Info (Regular, 7.5pt, Georgia)
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    contact.paragraph_format.space_before = Pt(4)
    contact.paragraph_format.space_after = Pt(0)
    
    contact_text = f"{personal['location']} | {personal['email']} | {personal['phone']} | "
    contact_run = contact.add_run(contact_text)
    contact_run.font.name = 'Georgia'
    contact_run.font.size = Pt(7.5)
    
    # Add hyperlinks for LinkedIn, GitHub, Website
    add_hyperlink(contact, personal['linkedin'], f"https://{personal['linkedin']}")
    pipe1 = contact.add_run(" | ")
    pipe1.font.name = 'Georgia'
    pipe1.font.size = Pt(7.5)
    
    add_hyperlink(contact, personal['github'], f"https://{personal['github']}")
    pipe2 = contact.add_run(" | ")
    pipe2.font.name = 'Georgia'
    pipe2.font.size = Pt(7.5)
    
    add_hyperlink(contact, personal['website'], f"https://{personal['website']}")
    
    # Add section divider after header
    add_horizontal_line(doc)
    
    # ===========================
    # SUMMARY SECTION
    # ===========================
    summary_items = filter_by_tags(data["summary"], role)
    
    if summary_items:
        # Section Header (Bold, 8pt, Georgia)
        summary_header = doc.add_paragraph()
        summary_header_run = summary_header.add_run("SUMMARY")
        summary_header_run.bold = True
        summary_header_run.font.name = 'Georgia'
        summary_header_run.font.size = Pt(8)
        summary_header.paragraph_format.space_after = Pt(0)
        summary_header.paragraph_format.space_before = Pt(6)
        
        # Summary text as continuous paragraph (7.5pt, Georgia)
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_before = Pt(4)
        summary_para.paragraph_format.space_after = Pt(0)
        summary_para.paragraph_format.line_spacing = 1.0
        
        summary_text = " ".join([item["text"] for item in summary_items])
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.name = 'Georgia'
        summary_run.font.size = Pt(7.5)
        
        # Add section divider after summary
        add_horizontal_line(doc)
    
    # ===========================
    # EXPERIENCE SECTION
    # ===========================
    experience_items = []
    for exp in data["experience"]:
        filtered_bullets = filter_by_tags(exp["bullets"], role)
        if filtered_bullets:
            experience_items.append({
                "company": exp["company"],
                "role": exp["role"],
                "location": exp["location"],
                "dates": exp["dates"],
                "bullets": filtered_bullets
            })
    
    if experience_items:
        # Section Header (Bold, 8pt, Georgia)
        exp_header = doc.add_paragraph()
        exp_header_run = exp_header.add_run("EXPERIENCE")
        exp_header_run.bold = True
        exp_header_run.font.name = 'Georgia'
        exp_header_run.font.size = Pt(8)
        exp_header.paragraph_format.space_after = Pt(0)
        exp_header.paragraph_format.space_before = Pt(6)
        
        # Add section divider immediately after header
        add_horizontal_line(doc)
        
        for exp in experience_items:
            # Job Title Line with RIGHT-ALIGNED dates
            job_para = doc.add_paragraph()
            job_para.paragraph_format.space_before = Pt(4)
            job_para.paragraph_format.space_after = Pt(0)
            
            # Add role and company on left
            job_run = job_para.add_run(f"{exp['role']} | {exp['company']}")
            job_run.bold = True
            job_run.font.name = 'Georgia'
            job_run.font.size = Pt(7.5)
            
            location_run = job_para.add_run(f" ({exp['location']})")
            location_run.font.name = 'Georgia'
            location_run.font.size = Pt(7.5)
            
            # Add tab to push dates to the right
            job_para.add_run('\t')
            
            dates_run = job_para.add_run(exp['dates'])
            dates_run.bold = True
            dates_run.font.name = 'Georgia'
            dates_run.font.size = Pt(7.5)
            
            # Set right-aligned tab stop at right margin
            tab_stops = job_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Bullets (7.5pt, Georgia)
            for bullet in exp["bullets"]:
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.space_before = Pt(4)
                bullet_para.paragraph_format.space_after = Pt(0)
                bullet_para.paragraph_format.line_spacing = 1.0
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                apply_bold_to_text(bullet_para, bullet["text"])
    
    # ===========================
    # PROJECTS SECTION
    # ===========================
    project_items = []
    for proj in data["projects"]:
        filtered_bullets = filter_by_tags(proj["bullets"], role)
        if filtered_bullets:
            project_items.append({
                "name": proj["name"],
                "tech_stack": proj["tech_stack"],
                "dates": proj["dates"],
                "bullets": filtered_bullets
            })
    
    if project_items:
        # Section Header (Bold, 8pt, Georgia)
        proj_header = doc.add_paragraph()
        proj_header_run = proj_header.add_run("PROJECTS")
        proj_header_run.bold = True
        proj_header_run.font.name = 'Georgia'
        proj_header_run.font.size = Pt(8)
        proj_header.paragraph_format.space_after = Pt(0)
        proj_header.paragraph_format.space_before = Pt(6)
        
        # Add section divider immediately after header
        add_horizontal_line(doc)
        
        for proj in project_items:
            # Project Title Line with RIGHT-ALIGNED dates
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(4)
            proj_para.paragraph_format.space_after = Pt(0)
            
            name_run = proj_para.add_run(f"{proj['name']} | ")
            name_run.bold = True
            name_run.font.name = 'Georgia'
            name_run.font.size = Pt(7.5)
            
            tech_run = proj_para.add_run(proj['tech_stack'])
            tech_run.font.name = 'Georgia'
            tech_run.font.size = Pt(7.5)
            
            # Add tab to push dates to the right
            proj_para.add_run('\t')
            
            dates_run = proj_para.add_run(proj['dates'])
            dates_run.bold = True
            dates_run.font.name = 'Georgia'
            dates_run.font.size = Pt(7.5)
            
            # Set right-aligned tab stop at right margin
            tab_stops = proj_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Bullets (7.5pt, Georgia)
            for bullet in proj["bullets"]:
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.space_before = Pt(4)
                bullet_para.paragraph_format.space_after = Pt(0)
                bullet_para.paragraph_format.line_spacing = 1.0
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                apply_bold_to_text(bullet_para, bullet["text"])
    
    # ===========================
    # EDUCATION SECTION
    # ===========================
    education_items = filter_by_tags(data["education"], role)
    
    if education_items:
        # Section Header (Bold, 8pt, Georgia)
        edu_header = doc.add_paragraph()
        edu_header_run = edu_header.add_run("EDUCATION")
        edu_header_run.bold = True
        edu_header_run.font.name = 'Georgia'
        edu_header_run.font.size = Pt(8)
        edu_header.paragraph_format.space_after = Pt(0)
        edu_header.paragraph_format.space_before = Pt(6)
        
        # Add section divider immediately after header
        add_horizontal_line(doc)
        
        for edu in education_items:
            # Degree Line with RIGHT-ALIGNED dates (Bold, 7.5pt, Georgia)
            degree_para = doc.add_paragraph()
            degree_para.paragraph_format.space_before = Pt(4)
            degree_para.paragraph_format.space_after = Pt(0)
            
            degree_run = degree_para.add_run(f"{edu['degree']}")
            degree_run.bold = True
            degree_run.font.name = 'Georgia'
            degree_run.font.size = Pt(7.5)
            
            # Add tab to push dates to the right
            degree_para.add_run('\t')
            
            dates_run = degree_para.add_run(edu['dates'])
            dates_run.bold = True
            dates_run.font.name = 'Georgia'
            dates_run.font.size = Pt(7.5)
            
            # Set right-aligned tab stop at right margin
            tab_stops = degree_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Institution Line (Regular, 7.5pt, Georgia)
            inst_para = doc.add_paragraph()
            inst_para.paragraph_format.space_before = Pt(4)
            inst_para.paragraph_format.space_after = Pt(0)
            
            inst_run = inst_para.add_run(f"{edu['institution']} ({edu['location']})")
            inst_run.font.name = 'Georgia'
            inst_run.font.size = Pt(7.5)
    
    # ===========================
    # SKILLS SECTION
    # ===========================
    skills_data = data["skills"]
    skill_items = []
    
    for key, skill_cat in skills_data.items():
        tags = skill_cat.get("tags", [])
        if "all" in tags or role in tags:
            skill_items.append(skill_cat)
    
    if skill_items:
        # Section Header (Bold, 8pt, Georgia)
        skills_header = doc.add_paragraph()
        skills_header_run = skills_header.add_run("SKILLS")
        skills_header_run.bold = True
        skills_header_run.font.name = 'Georgia'
        skills_header_run.font.size = Pt(8)
        skills_header.paragraph_format.space_after = Pt(0)
        skills_header.paragraph_format.space_before = Pt(6)
        
        # Add section divider immediately after header
        add_horizontal_line(doc)
        
        for skill_cat in skill_items:
            skill_para = doc.add_paragraph()
            skill_para.paragraph_format.space_before = Pt(4)
            skill_para.paragraph_format.space_after = Pt(0)
            skill_para.paragraph_format.line_spacing = 1.0
            
            label_run = skill_para.add_run(f"{skill_cat['label']}: ")
            label_run.bold = True
            label_run.font.name = 'Georgia'
            label_run.font.size = Pt(7.5)
            
            items_run = skill_para.add_run(", ".join(skill_cat['items']))
            items_run.font.name = 'Georgia'
            items_run.font.size = Pt(7.5)
    
    # ===========================
    # SAVE DOCUMENT WITH VERSIONING
    # ===========================
    os.makedirs("outputs", exist_ok=True)
    
    # Find the next available version number
    version = 1
    while True:
        output_path = f"outputs/resume_{role}_{version}.docx"
        if not os.path.exists(output_path):
            break
        version += 1
    
    doc.save(output_path)
    
    print(f"✅ Generated resume for role: {role}")
    print(f"📄 Output: {output_path}")
    print(f"🔢 Version: {version}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python build_docx.py <role>")
        print("Example: python build_docx.py python")
        print("Available roles: python, java, fullstack, backend, frontend, cloud, devops, mobile, all")
        sys.exit(1)
    
    role = sys.argv[1].lower()
    build_resume(role)