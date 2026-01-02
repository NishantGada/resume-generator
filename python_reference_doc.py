from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING

# 1️⃣ Create new document
doc = Document()

font_style = "georgia"

# 2️⃣ Configure page margins (1 inch all sides)
section = doc.sections[0]
section.top_margin = Inches(0.4)
section.bottom_margin = Inches(0.4)
section.left_margin = Inches(0.4)
section.right_margin = Inches(0.4)

# 3️⃣ Configure Normal style (body text)
normal_style = doc.styles['Normal']
normal_style.font.name = font_style
normal_style.font.size = Pt(7.5)
normal_style.font.color.rgb = RGBColor(0, 0, 0) # blue

# Add a placeholder paragraph to anchor formatting
p = doc.add_paragraph('Sample normal text')
p_format = p.paragraph_format
p_format.line_spacing = 0.6
p_format.space_before = Pt(0)
p_format.space_after = Pt(0)

# 4️⃣ Configure Heading 1 style (section titles)
# heading1 = doc.styles['Heading 1']
# heading1.font.name = font_style
# heading1.font.size = Pt(14)
# heading1.font.bold = True
# heading1.font.color.rgb = RGBColor(255, 0, 0) # red
# heading1.paragraph_format.line_spacing = 1.15
# heading1.paragraph_format.space_before = Pt(6)
# heading1.paragraph_format.space_after = Pt(3)

# Add a placeholder heading
doc.add_paragraph('Heading 1 Example', style='Heading 1')

# 5️⃣ Configure Heading 2 style (sub-section titles)
# heading2 = doc.styles['Heading 2']
# heading2.font.name = font_style
# heading2.font.size = Pt(12)
# heading2.font.bold = True
# heading2.font.color.rgb = RGBColor(0, 0, 255) # blue
# heading2.paragraph_format.line_spacing = 1.15
# heading2.paragraph_format.space_before = Pt(4)
# heading2.paragraph_format.space_after = Pt(2)


# 5️⃣ Configure Heading 3 style (sub-section titles)
heading3 = doc.styles['Heading 3']
heading3.font.name = 'georgia'
heading3.font.size = Pt(11)
heading3.font.bold = True
heading3.font.color.rgb = RGBColor(128, 0, 128)
heading3.paragraph_format.line_spacing = 1.15
heading3.paragraph_format.space_before = Pt(4)
heading3.paragraph_format.space_after = Pt(0)

# Add a placeholder sub-heading
doc.add_paragraph('Heading 2 Example', style='Heading 2')

# 6️⃣ Configure List Bullet style
bullet_style = doc.styles['List Bullet']
bullet_style.font.name = font_style
bullet_style.font.size = Pt(11)
bullet_style.paragraph_format.line_spacing = 1.15
bullet_style.paragraph_format.space_before = Pt(0)
bullet_style.paragraph_format.space_after = Pt(0)

# Add a sample bullet
doc.add_paragraph('Sample bullet point', style='List Bullet')

# 7️⃣ Save the reference DOCX
doc.save('reference.docx')

print("✅ ATS-friendly reference.docx generated successfully!")
